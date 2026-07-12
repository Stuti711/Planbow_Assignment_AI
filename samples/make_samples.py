"""Generate sample documents for end-to-end testing into samples/out/.

Includes deliberate defects so the validation layer provably fires:
  - invoice.pdf: printed total (1,150.00) != subtotal + tax (1,100.00)
  - resume.pdf:  malformed email 'priya.sharma[at]example.com'
random_notes.txt should classify as 'unknown'.
"""
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

OUT = Path(__file__).resolve().parent / "out"
OUT.mkdir(exist_ok=True)


def write_pdf(name: str, lines: list[tuple[str, int, bool]]) -> None:
    """lines: (text, font_size, bold)."""
    c = canvas.Canvas(str(OUT / name), pagesize=LETTER)
    _, height = LETTER
    y = height - 60
    for text, size, bold in lines:
        if y < 60:
            c.showPage()
            y = height - 60
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        c.drawString(60, y, text)
        y -= size + 8
    c.save()
    print(f"wrote {OUT / name}")


def t(text: str, size: int = 10, bold: bool = False) -> tuple[str, int, bool]:
    return (text, size, bold)


# --------------------------------------------------------------------------- invoice
write_pdf("invoice.pdf", [
    t("INVOICE", 18, True),
    t(""),
    t("Northwind Supplies Pvt Ltd", 12, True),
    t("42 Industrial Estate, Pune, MH 411001, India"),
    t(""),
    t("Bill To: Acme Software Solutions"),
    t("221B Tech Park, Bengaluru, KA 560103, India"),
    t(""),
    t("Invoice Number: INV-2026-0417", 11, True),
    t("Invoice Date: 2026-06-28"),
    t("Due Date: 2026-07-28"),
    t("Currency: INR"),
    t(""),
    t("Description                       Qty     Unit Price     Amount", 10, True),
    t("Ergonomic office chair              4        12,500.00    50,000.00"),
    t("Standing desk (electric)            2        22,500.00    45,000.00"),
    t("Monitor arm (dual)                  2         2,500.00     5,000.00"),
    t(""),
    t("Subtotal:                                                100,000.00"),
    t("GST (10%):                                                10,000.00"),
    # Deliberate defect: correct total is 110,000.00
    t("TOTAL DUE:                                               115,000.00", 12, True),
    t(""),
    t("Payment Terms: Net 30. Please quote the invoice number on payment."),
])

# --------------------------------------------------------------------- purchase order
write_pdf("purchase_order.pdf", [
    t("PURCHASE ORDER", 18, True),
    t(""),
    t("PO Number: PO-88121", 11, True),
    t("Order Date: 2026-07-01"),
    t("Expected Delivery: 2026-07-15"),
    t("Currency: USD"),
    t(""),
    t("Buyer: Globex Manufacturing Inc."),
    t("1500 Commerce Blvd, Austin, TX 78701, USA"),
    t(""),
    t("Supplier: Falcon Components LLC"),
    t("77 Harbor Way, San Diego, CA 92101, USA"),
    t(""),
    t("Description                       Qty     Unit Price     Amount", 10, True),
    t("Aluminium bracket A-113           500          4.20      2,100.00"),
    t("Stainless bolt M8 (box/100)       120         11.50      1,380.00"),
    t("Rubber gasket 60mm                300          1.40        420.00"),
    t(""),
    t("Subtotal:                                                 3,900.00"),
    t("Sales Tax (8.25%):                                          321.75"),
    t("TOTAL:                                                    4,221.75", 12, True),
    t(""),
    t("Shipping Terms: FOB Destination. Deliver to the Austin plant, dock 4."),
])

# -------------------------------------------------------------------------- contract
write_pdf("contract.pdf", [
    t("SERVICE AGREEMENT", 18, True),
    t(""),
    t("This Service Agreement (the 'Agreement') is entered into as of 2026-05-01"),
    t("(the 'Effective Date') by and between:"),
    t(""),
    t("1. Meridian Analytics Ltd, a company registered in England ('Service Provider'); and"),
    t("2. Blue Harbor Retail Group PLC ('Client')."),
    t(""),
    t("1. SERVICES", 12, True),
    t("The Service Provider shall provide data analytics and reporting services"),
    t("as described in Schedule A."),
    t(""),
    t("2. TERM", 12, True),
    t("This Agreement commences on the Effective Date and continues until"),
    t("2028-04-30, unless terminated earlier under clause 5. The Agreement does"),
    t("not renew automatically."),
    t(""),
    t("3. FEES", 12, True),
    t("The Client shall pay fees of GBP 120,000.00 per annum, invoiced quarterly"),
    t("in arrears, payable within 30 days of invoice."),
    t(""),
    t("4. GOVERNING LAW", 12, True),
    t("This Agreement is governed by the laws of England and Wales."),
    t(""),
    t("5. TERMINATION", 12, True),
    t("Either party may terminate with 90 days written notice, or immediately"),
    t("upon material breach that remains uncured for 30 days."),
])

# ---------------------------------------------------------------------------- resume
write_pdf("resume.pdf", [
    t("PRIYA SHARMA", 18, True),
    t("Senior Backend Engineer"),
    t(""),
    # Deliberate defect: malformed email
    t("Email: priya.sharma[at]example.com | Phone: +91 98765 43210"),
    t("Location: Hyderabad, India | GitHub: github.com/priyasharma-dev"),
    t(""),
    t("SUMMARY", 12, True),
    t("Backend engineer with 8 years of experience building distributed systems,"),
    t("payment platforms and data pipelines in Python and Go."),
    t(""),
    t("SKILLS", 12, True),
    t("Python, Go, PostgreSQL, Kafka, Docker, Kubernetes, AWS, gRPC, Redis"),
    t(""),
    t("EXPERIENCE", 12, True),
    t("Staff Engineer - FinEdge Payments | 2022-03 to Present", 11, True),
    t("Led the settlement platform team; cut reconciliation time by 70%."),
    t(""),
    t("Senior Software Engineer - CloudCart | 2018-06 to 2022-02", 11, True),
    t("Built order management microservices handling 40k orders/day."),
    t(""),
    t("Software Engineer - Zenlabs | 2016-07 to 2018-05", 11, True),
    t("Developed internal ETL tooling and REST APIs."),
    t(""),
    t("EDUCATION", 12, True),
    t("B.Tech in Computer Science, NIT Warangal, 2016"),
])

# ------------------------------------------------------------------------ resume DOCX
docx = DocxDocument()
docx.add_heading("ARJUN MEHTA", level=0)
docx.add_paragraph("Product Manager")
docx.add_paragraph("Email: arjun.mehta@example.com | Phone: +91 91234 56789 | Mumbai, India")
docx.add_heading("Summary", level=1)
docx.add_paragraph("Product manager with 6 years of experience shipping B2B SaaS "
                   "products from discovery to scale.")
docx.add_heading("Skills", level=1)
docx.add_paragraph("Roadmapping, SQL, A/B testing, Figma, Stakeholder management, Jira")
docx.add_heading("Experience", level=1)
docx.add_paragraph("Senior Product Manager — Planbow Technologies | 2023-01 to Present")
docx.add_paragraph("Own the collaboration suite; grew weekly active teams 3x.")
docx.add_paragraph("Product Manager — BrightDesk | 2019-04 to 2022-12")
docx.add_paragraph("Launched the analytics module adopted by 400+ customers.")
docx.add_heading("Education", level=1)
docx.add_paragraph("MBA, IIM Indore, 2019")
docx.save(str(OUT / "resume_arjun.docx"))
print(f"wrote {OUT / 'resume_arjun.docx'}")

# ------------------------------------------------------------------------- random TXT
(OUT / "random_notes.txt").write_text(
    "Reminder: water the plants on Tuesday.\n"
    "Grocery list: oats, almonds, coffee, bananas.\n"
    "Random thought: the pigeons on the balcony seem to prefer the left railing.\n"
    "Book to read next: something about typography, maybe.\n",
    encoding="utf-8",
)
print(f"wrote {OUT / 'random_notes.txt'}")

print("\nAll samples generated in", OUT)
