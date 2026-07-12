"""Deterministic ingestion: format detection, file saving, text extraction.

PDFs and images are passed to the model natively (no OCR pipeline);
DOCX and TXT are reduced to plain text here.
"""
import hashlib
import re
import uuid
from pathlib import Path
from typing import Union

from docx import Document as DocxDocument

from .config import settings

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

EXTENSION_MIME = {
    ".pdf": "application/pdf",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".docx": DOCX_MIME,
    ".txt": "text/plain",
}

# Magic-byte signatures for binary formats (DOCX is a zip, like any OOXML file).
MAGIC_SIGNATURES = {
    "application/pdf": b"%PDF",
    "image/png": b"\x89PNG",
    "image/jpeg": b"\xff\xd8",
    DOCX_MIME: b"PK\x03\x04",
}


class UnsupportedFormatError(ValueError):
    pass


def detect_mime(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    mime = EXTENSION_MIME.get(ext)
    if mime is None:
        raise UnsupportedFormatError(
            f"Unsupported file format '{ext}'. Supported: {', '.join(sorted(EXTENSION_MIME))}"
        )
    signature = MAGIC_SIGNATURES.get(mime)
    if signature is not None and not data.startswith(signature):
        raise UnsupportedFormatError(
            f"File content does not look like {mime} (extension/content mismatch)."
        )
    return mime


def save_upload(filename: str, data: bytes) -> tuple[str, str, str]:
    """Save an upload; returns (stored_path, sha256, mime_type)."""
    mime = detect_mime(filename, data)
    digest = hashlib.sha256(data).hexdigest()
    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", Path(filename).name)
    stored = settings.upload_dir / f"{uuid.uuid4().hex[:8]}_{safe_name}"
    stored.write_bytes(data)
    return str(stored), digest, mime


def _docx_to_text(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def load_content(file_path: str, mime_type: str) -> tuple[str, Union[bytes, str]]:
    """Load a stored file as model input.

    Returns ("binary", bytes) for PDF/images (sent to Gemini natively)
    or ("text", str) for DOCX/TXT.
    """
    path = Path(file_path)
    if mime_type in ("application/pdf", "image/png", "image/jpeg"):
        return "binary", path.read_bytes()
    if mime_type == DOCX_MIME:
        return "text", _docx_to_text(path)
    if mime_type == "text/plain":
        try:
            return "text", path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return "text", path.read_text(encoding="latin-1")
    raise UnsupportedFormatError(f"Unsupported mime type: {mime_type}")
